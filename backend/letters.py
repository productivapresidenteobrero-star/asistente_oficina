import os
import logging
from datetime import datetime
from pathlib import Path
from backend.config import CARTAS_PATH, COMUNA_INFO, OFICIO_PREFIJO
from backend.database import get_db, obtener_siguiente_numero_oficio
from backend.ai_router import generar_texto

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("letters")

CARTAS_PATH.mkdir(parents=True, exist_ok=True)

# Plantilla HTML con estilos CSS modernos y limpios para la carta oficial
PLANTILLA_CARTA_HTML = """
<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <title>Oficio Comunal</title>
    <style>
        @page {{
            size: letter;
            margin: 3cm 2cm 2.5cm 2cm;
            @top-center {{
                content: element(header);
            }}
            @bottom-center {{
                content: element(footer);
            }}
        }}
        
        body {{
            font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif;
            color: #333333;
            font-size: 11pt;
            line-height: 1.6;
        }}
        
        div.header {{
            position: running(header);
            border-bottom: 2px solid #1a365d;
            padding-bottom: 10px;
            margin-bottom: 20px;
        }}
        
        .header-table {{
            width: 100%;
            border-collapse: collapse;
        }}
        
        .header-title {{
            font-size: 14pt;
            font-weight: bold;
            color: #1a365d;
            text-transform: uppercase;
        }}
        
        .header-sub {{
            font-size: 9pt;
            color: #4a5568;
        }}
        
        div.footer {{
            position: running(footer);
            border-top: 1px solid #e2e8f0;
            padding-top: 5px;
            font-size: 8pt;
            color: #718096;
            text-align: center;
            width: 100%;
        }}
        
        .oficio-meta {{
            text-align: right;
            margin-bottom: 30px;
            font-size: 10pt;
        }}
        
        .oficio-numero {{
            font-weight: bold;
            color: #c53030;
            font-size: 12pt;
        }}
        
        .destinatario {{
            margin-bottom: 25px;
            font-weight: bold;
        }}
        
        .destinatario p {{
            margin: 2px 0;
        }}
        
        .asunto {{
            margin-bottom: 25px;
            font-weight: bold;
            text-transform: uppercase;
            border-bottom: 1px dashed #cbd5e0;
            padding-bottom: 5px;
        }}
        
        .contenido {{
            text-align: justify;
            margin-bottom: 50px;
            white-space: pre-line;
        }}
        
        .firma-area {{
            margin-top: 50px;
            text-align: center;
            page-break-inside: avoid;
        }}
        
        .firma-linea {{
            width: 250px;
            border-top: 1px solid #000;
            margin: 0 auto 10px auto;
        }}
        
        .firma-nombre {{
            font-weight: bold;
            font-size: 10pt;
        }}
        
        .firma-cargo {{
            font-size: 9pt;
            color: #4a5568;
        }}
    </style>
</head>
<body>

    <div class="header">
        <table class="header-table">
            <tr>
                <td>
                    <div class="header-title">{comuna_nombre}</div>
                    <div class="header-sub">{comuna_ubicacion} | RIF: {comuna_rif}</div>
                </td>
            </tr>
        </table>
    </div>

    <div class="footer">
        "La Patria es el hombre, el trabajo, el estudio y la moral." - Correo: {comuna_email}
    </div>

    <div class="oficio-meta">
        <p><strong>Fecha:</strong> {fecha}</p>
        <p><strong>Oficio Nro:</strong> <span class="oficio-numero">{numero_oficio}</span></p>
    </div>

    <div class="destinatario">
        <p>Ciudadano(a):</p>
        <p>{destinatario}</p>
        <p>{cargo_destinatario}</p>
        <p>Presente.-</p>
    </div>

    <div class="asunto">
        <strong>Asunto:</strong> {asunto}
    </div>

    <div class="contenido">
        {contenido}
    </div>

    <div class="firma-area">
        <div class="firma-linea"></div>
        <div class="firma-nombre">{firma_nombre}</div>
        <div class="firma-cargo">{firma_cargo}</div>
        <div class="firma-cargo">{firma_comuna}</div>
    </div>

</body>
</html>
"""

async def redactar_carta_ai(tipo: str, destinatario: str, asunto: str, instrucciones: str) -> str:
    """Usa el router de IA para redactar el cuerpo de la carta."""
    prompt = f"""
Genera el contenido de una carta formal de tipo '{tipo}' redactada en español.

Destinatario: {destinatario}
Asunto: {asunto}
Instrucciones específicas del contenido: {instrucciones}

Requisitos adicionales:
1. Usa un tono formal, sumamente educado y respetuoso, adecuado para la comunicación comunitaria de la Comuna Productiva Presidente Obrero en Venezuela.
2. Estructura el cuerpo con un saludo formal, una introducción de la situación, el desarrollo de la petición o mensaje y una despedida formal.
3. No incluyas el encabezado de fecha o membrete, ni firmas finales repetidas, ya que la plantilla HTML los colocará automáticamente. Solo redacta el cuerpo del mensaje y el párrafo de despedida.
"""
    system_instruction = "Eres un redactor experto en correspondencia comunitaria y oficios públicos formales en Venezuela."
    
    return await generar_texto(prompt, system_instruction=system_instruction)

async def html_a_pdf(html_content: str, pdf_filepath) -> bool:
    """Convierte HTML a PDF usando Playwright (Chromium headless)."""
    try:
        from playwright.async_api import async_playwright
        async with async_playwright() as p:
            browser = await p.chromium.launch()
            page = await browser.new_page()
            await page.set_content(html_content, wait_until="networkidle")
            await page.pdf(
                path=str(pdf_filepath),
                format="Letter",
                margin={"top": "2.5cm", "bottom": "2.5cm", "left": "2cm", "right": "2cm"},
                print_background=True
            )
            await browser.close()
        logger.info(f"PDF generado correctamente en: {pdf_filepath}")
        return True
    except Exception as e:
        logger.error(f"Error al generar PDF con Playwright: {e}")
        return False

def carta_a_docx(numero_oficio: str, fecha: str, destinatario: str, 
                 cargo_destinatario: str, asunto: str, contenido: str, 
                 docx_filepath, firma_nombre: str = "Enrique Maduro", 
                 firma_cargo: str = "Responsable de la Sala de Autogobierno") -> bool:
    """Genera una carta en formato .docx usando python-docx."""
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        
        # Márgenes de página
        section = doc.sections[0]
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

        # Encabezado
        encabezado = doc.add_paragraph()
        run_titulo = encabezado.add_run(COMUNA_INFO["nombre"].upper())
        run_titulo.bold = True
        run_titulo.font.size = Pt(14)
        encabezado.alignment = WD_ALIGN_PARAGRAPH.CENTER

        ubicacion = f"Edo. {COMUNA_INFO['estado']}"
        if COMUNA_INFO.get('municipio'):
            ubicacion = f"Mun. {COMUNA_INFO['municipio']}, {ubicacion}"
        subtitulo = doc.add_paragraph(f"{ubicacion} | RIF: {COMUNA_INFO.get('rif', 'G-20000000-0')}")
        subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        # Metadatos del oficio
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        meta.add_run(f"Fecha: {fecha}\n").bold = False
        run_nro = meta.add_run(f"Oficio Nro: {numero_oficio}")
        run_nro.bold = True

        doc.add_paragraph()

        # Destinatario
        dest_p = doc.add_paragraph()
        dest_p.add_run("Ciudadano(a):\n").bold = True
        dest_p.add_run(f"{destinatario}\n{cargo_destinatario or 'Ciudadano(a)'}\nPresente.-")

        doc.add_paragraph()

        # Asunto
        asunto_p = doc.add_paragraph()
        asunto_p.add_run("Asunto: ").bold = True
        asunto_p.add_run(asunto.upper())

        doc.add_paragraph()

        # Contenido
        for linea in contenido.split("\n"):
            p = doc.add_paragraph(linea)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        doc.add_paragraph()
        doc.add_paragraph()

        # Firma
        firma_p = doc.add_paragraph("_" * 35)
        firma_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cargo_p = doc.add_paragraph(firma_nombre)
        cargo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cargo_p.runs[0].bold = True
        cargo_p2 = doc.add_paragraph(firma_cargo)
        cargo_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        com_p = doc.add_paragraph(COMUNA_INFO["nombre"])
        com_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.save(str(docx_filepath))
        logger.info(f"DOCX generado correctamente en: {docx_filepath}")
        return True
    except Exception as e:
        logger.error(f"Error al generar DOCX: {e}")
        return False


async def crear_y_guardar_carta(
    tipo: str, 
    destinatario: str, 
    cargo_destinatario: str, 
    asunto: str, 
    instrucciones: str,
    vocero_firma_id: int = None
) -> dict:
    """Genera, guarda en PDF/DOCX y registra en DB una carta comunal con número de oficio automático."""
    
    # 1. Obtener número de oficio CPPO
    numero_oficio = await obtener_siguiente_numero_oficio(prefijo=OFICIO_PREFIJO)
    
    # 2. Resolver vocero firmante (fijo: Enrique Maduro si no se selecciona otro)
    vocero_nombre = "Enrique Maduro"
    vocero_cargo = "Responsable de la Sala de Autogobierno"
    vocero_telefono = "04166880411"
    
    if vocero_firma_id:
        async with await get_db() as db:
            cursor = await db.execute("SELECT nombre, cargo, telefono FROM voceros WHERE id=?", (vocero_firma_id,))
            row = await cursor.fetchone()
            if row:
                vocero_nombre = row[0] or vocero_nombre
                vocero_cargo = row[1] or vocero_cargo
                vocero_telefono = row[2] or vocero_telefono
    
    # 3. Redactar el cuerpo con IA
    contenido = await redactar_carta_ai(tipo, destinatario, asunto, instrucciones)
    
    # 4. Formatear la plantilla HTML
    fecha_actual_legible = datetime.now().strftime("%d de %B de %Y")
    meses = {
        "January": "Enero", "February": "Febrero", "March": "Marzo", "April": "Abril",
        "May": "Mayo", "June": "Junio", "July": "Julio", "August": "Agosto",
        "September": "Septiembre", "October": "Octubre", "November": "Noviembre", "December": "Diciembre"
    }
    for eng, esp in meses.items():
        fecha_actual_legible = fecha_actual_legible.replace(eng, esp)
        
    ubicacion = f"Edo. {COMUNA_INFO['estado']}"
    if COMUNA_INFO['municipio']:
        ubicacion = f"Mun. {COMUNA_INFO['municipio']}, {ubicacion}"
        
    html_content = PLANTILLA_CARTA_HTML.format(
        comuna_nombre=COMUNA_INFO["nombre"],
        comuna_ubicacion=ubicacion,
        comuna_rif=COMUNA_INFO["rif"] or "G-20000000-0",
        comuna_email=COMUNA_INFO["email"],
        fecha=fecha_actual_legible,
        numero_oficio=numero_oficio,
        destinatario=destinatario,
        cargo_destinatario=cargo_destinatario or "Ciudadano(a)",
        asunto=asunto,
        contenido=contenido,
        firma_nombre=vocero_nombre,
        firma_cargo=vocero_cargo,
        firma_comuna=COMUNA_INFO["nombre"]
    )
    
    # 5. Guardar archivos: HTML fuente, PDF real, DOCX
    nombre_base = numero_oficio.replace('.', '-').replace('/', '-')
    html_filepath = CARTAS_PATH / f"{nombre_base}.html"
    pdf_filepath  = CARTAS_PATH / f"{nombre_base}.pdf"
    docx_filepath = CARTAS_PATH / f"{nombre_base}.docx"
    
    # HTML (fuente de respaldo)
    try:
        with open(html_filepath, 'w', encoding='utf-8') as f:
            f.write(html_content)
    except Exception as e:
        logger.error(f"Error al guardar HTML fuente: {e}")

    # PDF mediante Playwright
    pdf_ok = await html_a_pdf(html_content, pdf_filepath)
    pdf_path_str = str(pdf_filepath.resolve()) if pdf_ok else str(html_filepath.resolve())

    # DOCX mediante python-docx
    carta_a_docx(
        numero_oficio=numero_oficio,
        fecha=fecha_actual_legible,
        destinatario=destinatario,
        cargo_destinatario=cargo_destinatario,
        asunto=asunto,
        contenido=contenido,
        docx_filepath=docx_filepath,
        firma_nombre=vocero_nombre,
        firma_cargo=vocero_cargo
    )
    docx_path_str = str(docx_filepath.resolve())

    # 6. Guardar metadatos en SQLite
    async with await get_db() as db:
        destinatario_db = f"{destinatario} ({cargo_destinatario})" if cargo_destinatario else destinatario
        await db.execute(
            """
            INSERT INTO cartas (numero_oficio, tipo, fecha, destinatario, asunto, contenido, pdf_path, docx_path, vocero_firma)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                numero_oficio, 
                tipo, 
                datetime.now().strftime("%Y-%m-%d"), 
                destinatario_db, 
                asunto, 
                contenido, 
                pdf_path_str,
                docx_path_str,
                f"{vocero_nombre} - {vocero_cargo}"
            )
        )
        await db.commit()
        
    return {
        "numero_oficio": numero_oficio,
        "tipo": tipo,
        "destinatario": destinatario,
        "asunto": asunto,
        "contenido": contenido,
        "pdf_path": pdf_path_str,
        "docx_path": docx_path_str,
        "vocero_firma": f"{vocero_nombre} - {vocero_cargo}"
    }
